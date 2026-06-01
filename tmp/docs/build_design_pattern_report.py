from pathlib import Path
import math
import shutil
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"E:\JavaProject\sell")
SRC = Path(r"D:\chrome download\设计模式-实验报告 (1).docx")
OUT = ROOT / "output" / "doc" / "设计模式-实验报告-完成版.docx"
TMP = ROOT / "tmp" / "docs"

FONT_CN = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_CN_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_LATIN = Path(r"C:\Windows\Fonts\arial.ttf")


def font(size: int, bold: bool = False):
    path = FONT_CN_BOLD if bold and FONT_CN_BOLD.exists() else FONT_CN
    return ImageFont.truetype(str(path), size=size)


def wrap_text(text, draw, fnt, max_width):
    lines = []
    for raw in text.split("\n"):
        if not raw:
            lines.append("")
            continue
        current = ""
        for ch in raw:
            test = current + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_box(draw, xy, title, lines, fill="#ffffff", outline="#2f5597"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=3)
    title_font = font(25, True)
    body_font = font(21)
    title_lines = title.split("\n")
    header_h = 42 + 28 * (len(title_lines) - 1)
    draw.rounded_rectangle((x1, y1, x2, y1 + header_h), radius=10, fill="#dbeafe", outline=outline, width=0)
    draw.line((x1, y1 + header_h, x2, y1 + header_h), fill=outline, width=2)
    ty = y1 + 7
    for title_line in title_lines:
        tw = draw.textbbox((0, 0), title_line, font=title_font)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, ty), title_line, fill="#17365d", font=title_font)
        ty += 28
    y = y1 + header_h + 10
    for line in lines:
        draw.text((x1 + 14, y), line, fill="#1f2937", font=body_font)
        y += 28


def arrow(draw, start, end, label="", color="#334155", width=3, dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        steps = max(1, int(math.hypot(x2 - x1, y2 - y1) / 18))
        for i in range(steps):
            if i % 2 == 0:
                a = i / steps
                b = (i + 1) / steps
                draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill=color, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 13
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - 0.5), y2 - size * math.sin(ang - 0.5)),
        (x2 - size * math.cos(ang + 0.5), y2 - size * math.sin(ang + 0.5)),
    ]
    draw.polygon(pts, fill=color)
    if label:
        lf = font(20)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        bbox = draw.textbbox((0, 0), label, font=lf)
        draw.rectangle((mx - 8, my - 15, mx + bbox[2] + 8, my + 16), fill="#ffffff")
        draw.text((mx, my - 13), label, fill=color, font=lf)


def polyline_arrow(draw, points, label="", color="#334155", width=3, dashed=False):
    for start, end in zip(points, points[1:]):
        x1, y1 = start
        x2, y2 = end
        if dashed:
            steps = max(1, int(math.hypot(x2 - x1, y2 - y1) / 18))
            for i in range(steps):
                if i % 2 == 0:
                    a = i / steps
                    b = (i + 1) / steps
                    draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill=color, width=width)
        else:
            draw.line((x1, y1, x2, y2), fill=color, width=width)
    arrow(draw, points[-2], points[-1], "", color=color, width=width, dashed=False)
    if label:
        lf = font(20)
        lx, ly = points[len(points) // 2]
        bbox = draw.textbbox((0, 0), label, font=lf)
        draw.rectangle((lx - 8, ly - 15, lx + bbox[2] + 8, ly + 16), fill="#ffffff")
        draw.text((lx, ly - 13), label, fill=color, font=lf)


def draw_note(draw, xy, text):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=8, fill="#f1f5f9", outline="#cbd5e1", width=2)
    lines = wrap_text(text, draw, font(22), x2 - x1 - 34)
    y = y1 + 16
    for line in lines:
        draw.text((x1 + 18, y), line, fill="#334155", font=font(22))
        y += 32


def create_canvas(title, w=1500, h=900):
    img = Image.new("RGB", (w, h), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((42, 24), title, fill="#0f172a", font=font(34, True))
    draw.line((42, 72, w - 42, 72), fill="#94a3b8", width=2)
    return img, draw


def save_abstract_factory():
    img, d = create_canvas("类图1：抽象工厂模式 - FPS 游戏场景族创建", w=1900, h=1220)
    draw_box(d, (120, 130, 500, 340), "<<interface>>\nSceneFactory", ["+createMap(): Map", "+createWeather(): Weather", "+createSound(): Sound"], "#ffffff")
    draw_box(d, (80, 760, 440, 945), "DesertSceneFactory", ["+createMap(): Map", "+createWeather(): Weather", "+createSound(): Sound"], "#fff7ed")
    draw_box(d, (500, 760, 860, 945), "SnowSceneFactory", ["+createMap(): Map", "+createWeather(): Weather", "+createSound(): Sound"], "#eff6ff")
    draw_box(d, (920, 760, 1280, 945), "CitySceneFactory", ["+createMap(): Map", "+createWeather(): Weather", "+createSound(): Sound"], "#f0fdf4")
    arrow(d, (260, 760), (260, 341), dashed=True)
    arrow(d, (680, 760), (390, 341), dashed=True)
    arrow(d, (1100, 760), (475, 341), dashed=True)

    draw_box(d, (730, 105, 1080, 240), "<<interface>> Map", ["+load()"], "#ffffff")
    draw_box(d, (730, 285, 1080, 420), "<<interface>> Weather", ["+show()"], "#ffffff")
    draw_box(d, (730, 465, 1080, 600), "<<interface>> Sound", ["+play()"], "#ffffff")
    arrow(d, (500, 205), (730, 172), "creates")
    arrow(d, (500, 235), (730, 352), "creates")
    arrow(d, (500, 265), (730, 532), "creates")

    draw_box(d, (1390, 105, 1740, 240), "Map 产品族", ["DesertMap", "SnowMap", "CityMap"], "#f8fafc")
    draw_box(d, (1390, 345, 1740, 520), "Weather 产品族", ["SandstormWeather", "SnowWeather", "RainWeather"], "#f8fafc")
    draw_box(d, (1390, 625, 1740, 800), "Sound 产品族", ["DesertSound", "SnowSound", "CitySound"], "#f8fafc")
    polyline_arrow(d, [(1390, 172), (1081, 172)], dashed=True)
    polyline_arrow(d, [(1390, 432), (1230, 432), (1230, 352), (1081, 352)], dashed=True)
    polyline_arrow(d, [(1390, 712), (1230, 712), (1230, 532), (1081, 532)], dashed=True)
    draw_note(d, (120, 1010, 1740, 1090), "说明：每个具体场景工厂创建同一产品族中的 Map、Weather、Sound。新增场景时增加一个具体工厂和一组具体产品即可。")
    path = TMP / "diagram_abstract_factory.png"
    img.save(path)
    return path


def save_observer():
    img, d = create_canvas("类图2：观察者模式 - FPS 系统通知模块", w=1600, h=980)
    draw_box(d, (95, 130, 500, 340), "<<interface>>\nNotificationSubject", ["+attach(member)", "+detach(member)", "+notify(message)"], "#ffffff")
    draw_box(d, (120, 610, 520, 805), "GameNotificationCenter", ["-members: List<GameMember>", "+attach(member)", "+detach(member)", "+notify(message)"], "#e0f2fe")
    polyline_arrow(d, [(320, 610), (320, 341)], dashed=True)

    draw_box(d, (1010, 130, 1350, 315), "<<interface>>\nGameMember", ["+update(message)"], "#ffffff")
    draw_box(d, (770, 610, 1010, 800), "Soldier", ["-name: String", "+update(message)"], "#f0fdf4")
    draw_box(d, (1060, 610, 1300, 800), "Commander", ["-name: String", "+update(message)"], "#fff7ed")
    draw_box(d, (1350, 610, 1590, 800), "Sniper", ["-name: String", "+update(message)"], "#fef2f2")
    arrow(d, (520, 705), (770, 705), "notify 1..*")
    polyline_arrow(d, [(890, 610), (890, 470), (1110, 470), (1110, 316)], dashed=True)
    polyline_arrow(d, [(1180, 610), (1180, 316)], dashed=True)
    polyline_arrow(d, [(1470, 610), (1470, 470), (1250, 470), (1250, 316)], dashed=True)
    draw_note(d, (95, 850, 1505, 930), "说明：通知中心保存 GameMember 列表，系统通知到达时统一调用 update(message)。成员可在运行时注册或移除。")
    path = TMP / "diagram_observer.png"
    img.save(path)
    return path


def save_singleton_facade():
    img, d = create_canvas("类图3：单例模式 + 外观模式 - 游戏管理器", w=1600, h=1020)
    draw_box(d, (85, 390, 395, 570), "PlayerClient", ["+openManager()", "+saveSettings()"], "#ffffff")
    draw_box(d, (565, 170, 1015, 585), "GameManager", ["-instance: GameManager", "-soundSetting: SoundEffectSetting", "-sceneSetting: SceneSetting", "-roleSetting: RoleSetting", "+getInstance(): GameManager", "+setSoundEffect(params)", "+setScene(params)", "+setRole(params)", "+applyAll()"], "#e0f2fe")
    draw_box(d, (1180, 115, 1540, 285), "SoundEffectSetting", ["-volume: int", "-enable3D: boolean", "+configure(params)"], "#fff7ed")
    draw_box(d, (1180, 390, 1540, 560), "SceneSetting", ["-quality: String", "-brightness: int", "+configure(params)"], "#f0fdf4")
    draw_box(d, (1180, 665, 1540, 835), "RoleSetting", ["-sensitivity: int", "-skin: String", "+configure(params)"], "#fef2f2")
    arrow(d, (395, 480), (565, 375))
    arrow(d, (1015, 265), (1180, 200), "set")
    arrow(d, (1015, 380), (1180, 475), "set")
    arrow(d, (1015, 495), (1180, 750), "set")
    draw_note(d, (85, 880, 1540, 960), "说明：GameManager 使用单例模式保证界面唯一；同时作为外观类，为音效、场景、角色等子系统提供统一设置入口。")
    path = TMP / "diagram_singleton_facade.png"
    img.save(path)
    return path


def save_adapter_strategy():
    img, d = create_canvas("类图4：适配器模式 + 策略模式 - 场景渲染模块", w=1700, h=1000)
    draw_box(d, (80, 205, 425, 395), "SceneRenderer", ["-strategy: RenderStrategy", "+setStrategy(strategy)", "+render(scene)"], "#e0f2fe")
    draw_box(d, (650, 145, 1010, 330), "<<interface>>\nRenderStrategy", ["+render(scene)"], "#ffffff")
    arrow(d, (425, 300), (650, 238), "uses")

    draw_box(d, (420, 620, 720, 800), "SurfaceRendering", ["+render(scene)"], "#f0fdf4")
    draw_box(d, (790, 620, 1090, 800), "VolumeRendering", ["+render(scene)"], "#fff7ed")
    draw_box(d, (1160, 620, 1490, 800), "RenderEngineAdapter", ["-engine: RenderEngine", "+render(scene)"], "#fef2f2")
    polyline_arrow(d, [(570, 620), (570, 500), (760, 500), (760, 331)], dashed=True)
    polyline_arrow(d, [(940, 620), (940, 331)], dashed=True)
    polyline_arrow(d, [(1325, 620), (1325, 500), (990, 500), (990, 331)], dashed=True)

    draw_box(d, (1190, 150, 1535, 330), "RenderEngine", ["+engineRender(data)", "+setEngineOptions()"], "#ffffff")
    arrow(d, (1325, 620), (1360, 331), "adapts")
    draw_note(d, (80, 850, 1535, 930), "说明：SceneRenderer 通过策略模式切换不同渲染算法；RenderEngineAdapter 通过适配器模式复用已有 RenderEngine。")
    path = TMP / "diagram_adapter_strategy.png"
    img.save(path)
    return path


def set_cell_text(cell, text, font_size=11):
    cell.text = ""
    for idx, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)


def add_paragraph(cell, text="", size=10.5, bold=False, color=None, style=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if style == "heading":
        p.paragraph_format.space_before = Pt(7)
    return p


def add_bullets(cell, items):
    for item in items:
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("· " + item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.text = ""


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def fill_label_value_row(table, row_idx, value):
    cell = table.rows[row_idx].cells[1]
    set_cell_text(cell, value)
    for c in table.rows[row_idx].cells:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def build_doc():
    TMP.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    diagrams = [
        save_abstract_factory(),
        save_observer(),
        save_singleton_facade(),
        save_adapter_strategy(),
    ]

    doc = Document(str(SRC))
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    title = doc.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(18)
        run.bold = True

    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)

    fill_label_value_row(table, 2, "2026年5月13日")
    fill_label_value_row(table, 4, "Windows 11；Microsoft Word/WPS；UML 类图；面向对象设计方法")
    fill_label_value_row(table, 5, "设计模式在 FPS 游戏系统中的应用")
    fill_label_value_row(
        table,
        6,
        "① 学会选择合适的设计模式解决在软件开发中遇到的实际问题，加深对常用设计模式的理解和掌握。\n"
        "② 掌握一些常用的设计模式联用技巧，学会同时运用多种设计模式解决实际问题。",
    )
    fill_label_value_row(
        table,
        7,
        "本实验围绕某 FPS 游戏系统的场景创建、系统通知、游戏管理器和场景渲染四个模块展开设计。\n"
        "① 游戏场景模块：当用户选择某个场景时，同时创建该场景对应的 Map、Weather 和 Sound，并支持新增场景。\n"
        "② 通知模块：系统向所有游戏成员发送任务完成、新任务提醒、敌袭警报等通知，并支持灵活增删成员。\n"
        "③ 游戏管理器模块：用户通过唯一的 Game Manager 界面对 Sound effect、Scene、Role 等对象进行统一参数设置。\n"
        "④ 场景渲染模块：支持表面渲染、体渲染等可扩展算法，同时能够调用已有 Render Engine 中的渲染算法。",
    )

    process_cell = table.rows[8].cells[1]
    clear_cell(process_cell)
    add_paragraph(process_cell, "一、游戏场景模块设计", 12, True, "1f4e79", "heading")
    add_paragraph(process_cell, "采用设计模式：抽象工厂模式（Abstract Factory Pattern）。", bold=True)
    add_paragraph(process_cell, "选择理由：地图、天气和背景音乐属于同一游戏场景下的一组相关对象。抽象工厂可以把同一场景的对象创建过程封装在一个具体工厂中，使客户端只依赖抽象接口。新增场景时，只需增加新的具体工厂和对应具体产品类，符合开闭原则。")
    add_bullets(process_cell, [
        "SceneFactory：抽象工厂，声明 createMap()、createWeather()、createSound()。",
        "DesertSceneFactory、SnowSceneFactory、CitySceneFactory：具体工厂，分别创建沙漠、雪地、城市场景的一组对象。",
        "Map、Weather、Sound：抽象产品接口，定义 load()、show()、play() 等行为。",
        "DesertMap、SandstormWeather、DesertSound 等：具体产品，负责对应场景表现。",
    ])
    p = process_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagrams[0]), width=Cm(15.8))
    add_paragraph(process_cell, "图1  游戏场景模块抽象工厂模式类图", 9.5, False)

    add_paragraph(process_cell, "二、系统通知模块设计", 12, True, "1f4e79", "heading")
    add_paragraph(process_cell, "采用设计模式：观察者模式（Observer Pattern）。", bold=True)
    add_paragraph(process_cell, "选择理由：通知中心与游戏成员之间是一对多依赖关系。当系统产生通知时，所有已注册成员都应自动收到消息。观察者模式支持在运行时 attach 或 detach 成员，通知中心不需要知道具体成员类型，扩展性好。")
    add_bullets(process_cell, [
        "NotificationSubject：抽象主题，定义 attach()、detach()、notify()。",
        "GameNotificationCenter：具体主题，维护成员列表并广播消息。",
        "GameMember：抽象观察者，定义 update(message)。",
        "Soldier、Commander、Sniper：具体观察者，接收并处理通知。",
    ])
    p = process_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagrams[1]), width=Cm(15.8))
    add_paragraph(process_cell, "图2  系统通知模块观察者模式类图", 9.5, False)

    add_paragraph(process_cell, "三、游戏管理器模块设计", 12, True, "1f4e79", "heading")
    add_paragraph(process_cell, "采用设计模式：单例模式（Singleton Pattern）+ 外观模式（Facade Pattern）。", bold=True)
    add_paragraph(process_cell, "选择理由：单例模式保证游戏运行时只存在一个 GameManager 界面，避免多个设置窗口造成状态冲突和资源浪费；外观模式为音效、场景、角色等子系统提供统一入口，客户端通过 GameManager 即可完成多对象设置，降低调用复杂度。")
    add_bullets(process_cell, [
        "GameManager：单例外观类，保存唯一 instance，并对外提供 getInstance()、setSoundEffect()、setScene()、setRole()、applyAll()。",
        "SoundEffectSetting、SceneSetting、RoleSetting：子系统类，分别处理音效、场景、角色参数。",
        "PlayerClient：客户端，只与 GameManager 交互，不直接操作各子系统。",
    ])
    p = process_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagrams[2]), width=Cm(15.8))
    add_paragraph(process_cell, "图3  游戏管理器单例模式与外观模式类图", 9.5, False)

    add_paragraph(process_cell, "四、场景渲染模块设计", 12, True, "1f4e79", "heading")
    add_paragraph(process_cell, "采用设计模式：适配器模式（Adapter Pattern）+ 策略模式（Strategy Pattern）。", bold=True)
    add_paragraph(process_cell, "选择理由：策略模式把不同渲染算法封装为可替换的 RenderStrategy，便于在运行时切换表面渲染、体渲染或新增渲染效果；适配器模式把已有 RenderEngine 的接口转换为系统统一需要的 RenderStrategy 接口，从而复用已有引擎算法。")
    add_bullets(process_cell, [
        "SceneRenderer：上下文类，持有 RenderStrategy，可通过 setStrategy() 更换算法。",
        "RenderStrategy：策略接口，定义 render(scene)。",
        "SurfaceRendering、VolumeRendering：具体策略，分别实现表面渲染和体渲染。",
        "RenderEngine：已有渲染引擎，接口与本系统不一致。",
        "RenderEngineAdapter：适配器，实现 RenderStrategy，在 render(scene) 中调用已有引擎算法。",
    ])
    p = process_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagrams[3]), width=Cm(15.8))
    add_paragraph(process_cell, "图4  场景渲染模块适配器模式与策略模式类图", 9.5, False)

    result_cell = table.rows[9].cells[1]
    clear_cell(result_cell)
    add_paragraph(result_cell, "本实验完成了 FPS 游戏系统中四个典型模块的设计模式选型与结构设计：", bold=True)
    add_bullets(result_cell, [
        "场景创建模块使用抽象工厂模式，实现 Map、Weather、Sound 一组相关对象的统一创建。",
        "系统通知模块使用观察者模式，实现通知中心与游戏成员之间的松耦合消息广播。",
        "游戏管理器模块使用单例模式和外观模式，既保证管理器唯一，又统一封装多个设置子系统。",
        "场景渲染模块使用策略模式和适配器模式，既能灵活扩展渲染算法，又能复用已有渲染引擎。",
        "各模块均满足可扩展、低耦合和职责清晰的设计目标。",
    ])

    reflection_cell = table.rows[10].cells[1]
    clear_cell(reflection_cell)
    add_paragraph(reflection_cell, "通过本次实验，我进一步理解了设计模式并不是孤立使用的模板，而是面向具体问题的设计经验总结。抽象工厂适合创建同一产品族对象，观察者适合处理一对多通知，单例与外观组合可以同时解决唯一实例和统一入口问题，策略与适配器组合则兼顾算法扩展和旧系统复用。实验也说明，在软件设计中应先分析对象之间的变化点和依赖关系，再选择能够降低耦合、提高扩展性的模式。")

    for row_idx in [6, 7, 8, 9, 10]:
        shade_cell(table.rows[row_idx].cells[0], "D9EAF7")

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_doc()
